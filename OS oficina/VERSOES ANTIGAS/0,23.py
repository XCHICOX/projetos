import tkinter as tk
from tkinter import messagebox, ttk
from pathlib import Path
import os
import webbrowser
from docx import Document
from PIL import Image, ImageTk

# Função para carregar e redimensionar ícones
def carregar_icone(nome_icone, tamanho=(16, 16)):
    caminho_icone = icone_pasta / nome_icone
    imagem = Image.open(caminho_icone)
    imagem = imagem.resize(tamanho, Image.LANCZOS)  # Redimensiona a imagem usando Image.LANCZOS
    return ImageTk.PhotoImage(imagem)

# Função para obter o próximo número sequencial para a OS
def obter_numero_sequencial():
    arquivo_numero = Path('numero_sequencial.txt')
    
    if arquivo_numero.exists():
        with open(arquivo_numero, 'r') as file:
            numero = int(file.read().strip())
    else:
        numero = 0

    proximo_numero = numero + 1
    
    with open(arquivo_numero, 'w') as file:
        file.write(str(proximo_numero))
    
    return proximo_numero

# Função para substituir marcadores de posição no documento DOCX
def substituir_marcadores(doc, substituicoes):
    for par, valor in substituicoes.items():
        for paragrafo in doc.paragraphs:
            if par in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(par, valor)
            
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for par, valor in substituicoes.items():
                    if par in celula.text:
                        celula.text = celula.text.replace(par, valor)

# Função para salvar a ordem de serviço
def salvar_ordem():
    nome_cliente = entry_nome_cliente.get()
    telefone = entry_telefone.get()
    veiculo = entry_veiculo.get()

    # Verifica se todos os campos obrigatórios estão preenchidos
    if not nome_cliente or not telefone or not veiculo:
        messagebox.showerror("Erro", "Por favor, preencha todos os campos obrigatórios (Nome do Cliente, Telefone e Veículo).")
        return

    pecas = [e.get() for e in peca_entries]
    valores = [v.get() for v in valor_entries]
    quantidades = [q.get() for q in quantidade_entries]

    # Verifica se pelo menos uma linha de peças está preenchida
    if not any(peca and valor and quantidade for peca, valor, quantidade in zip(pecas, valores, quantidades)):
        messagebox.showerror("Erro", "Por favor, preencha pelo menos uma linha na aba 'Peças e Valores'.")
        return

    numero_sequencial = obter_numero_sequencial()
    nome_arquivo = f'{numero_sequencial:04d}.docx'  # Nome do arquivo apenas com o número

    # Carrega o modelo de documento
    doc = Document('modelo.docx')

    # Substitui os marcadores de posição no modelo
    substituicoes = {
        '{{NUMERO_OS}}': f'{numero_sequencial:04d}',
        '{{NOME_CLIENTE}}': nome_cliente if nome_cliente else '',
        '{{TELEFONE}}': telefone if telefone else '',
        '{{VEICULO}}': veiculo if veiculo else '',
        '{{PECAS_VALORES}}': ''  # Inicialmente vazio, será preenchido com a tabela
    }

    substituir_marcadores(doc, substituicoes)

    # Adiciona uma tabela para peças, quantidades, valores e totais
    tabela = doc.add_table(rows=1, cols=4)
    tabela.style = 'Table Grid'
    
    # Define o cabeçalho da tabela
    hdr_cells = tabela.rows[0].cells
    hdr_cells[0].text = 'Peças'
    hdr_cells[1].text = 'Valor Unitário'
    hdr_cells[2].text = 'Quantidade'
    hdr_cells[3].text = 'Total'

    total_geral = 0.0
    
    for peca, valor, quantidade in zip(pecas, valores, quantidades):
        try:
            valor_float = float(valor.replace(',', '.')) if valor else 0.0
            quantidade_int = int(quantidade) if quantidade else 0
            total_item = valor_float * quantidade_int
            total_geral += total_item
        except ValueError:
            continue  # Ignora valores inválidos

        row_cells = tabela.add_row().cells
        row_cells[0].text = peca if peca else ''
        row_cells[1].text = quantidade if quantidade else ''
        row_cells[2].text = valor if valor else ''
        row_cells[3].text = f'R$ {total_item:.2f}'

    # Adiciona uma linha de resumo para quantidade total e soma dos valores
    try:
        quantidade_total = sum(int(q) for q in quantidades if q.isdigit())
    except ValueError:
        quantidade_total = 0

    row_cells = tabela.add_row().cells
    row_cells[0].text = 'Resumo'
    row_cells[1].text = ''
    row_cells[2].text = ''
    row_cells[3].text = f'R$ {total_geral:.2f}'

    # Substitui o marcador de posição da tabela
    substituicoes['{{PECAS_VALORES}}'] = tabela

    # Salva o documento
    doc.save(nome_arquivo)
    messagebox.showinfo("Sucesso", f"Ordem de serviço salva com sucesso! ({nome_arquivo})")
    
    # Habilita o botão de abrir o documento
    botao_abrir.config(state=tk.NORMAL)
    arquivo_salvo[0] = nome_arquivo  # Atualiza o nome do arquivo salvo

# Função para abrir o arquivo salvo
def abrir_arquivo():
    if arquivo_salvo[0]:
        os.startfile(arquivo_salvo[0])
    else:
        messagebox.showerror("Erro", "Nenhum arquivo salvo para abrir.")

# Função para abrir o arquivo pelo número da OS
def abrir_arquivo_pelo_numero():
    numero_os_str = entry_numero_os.get()  # Obtém o número da OS como string
    try:
        numero_os = int(numero_os_str)  # Converte a string para inteiro
        nome_arquivo = f'{numero_os:04d}.docx'  # Formata o nome do arquivo com zeros à esquerda
        if Path(nome_arquivo).exists():
            os.startfile(nome_arquivo)  # Abre o arquivo
        else:
            messagebox.showerror("Erro", "Arquivo não encontrado.")
    except ValueError:
        messagebox.showerror("Erro", "Número da OS inválido. Por favor, insira um número válido.")

# Função para atualizar o total na interface
def atualizar_total(event=None):
    total_geral = 0.0
    for i in range(num_linhas):
        try:
            valor = float(valor_entries[i].get().replace(',', '.')) if valor_entries[i].get() else 0.0
            quantidade = int(quantidade_entries[i].get()) if quantidade_entries[i].get() else 0
            total = valor * quantidade
            total_entries[i].config(state='normal')
            total_entries[i].delete(0, tk.END)
            total_entries[i].insert(0, f'R$ {total:.2f}')
            total_entries[i].config(state='readonly')
            total_geral += total
        except ValueError:
            continue  # Ignora valores inválidos
    return total_geral

# Função para carregar documentos DOCX disponíveis
def carregar_docx():
    listbox_pdfs.delete(0, tk.END)
    pdfs = [f for f in Path('.').glob('*.docx')]
    for pdf in pdfs:
        listbox_pdfs.insert(tk.END, pdf.name)

# Função para abrir o arquivo selecionado na lista
def abrir_arquivo_selecionado(event):
    pdf_selecionado = listbox_pdfs.get(tk.ACTIVE)
    if pdf_selecionado:
        os.startfile(pdf_selecionado)
    else:
        messagebox.showerror("Erro", "Nenhum arquivo selecionado.")        

# Função para enviar mensagem pelo WhatsApp com o documento anexado
def enviar_whatsapp():
    pdf_selecionado = listbox_pdfs.get(tk.ACTIVE)
    if not pdf_selecionado:
        messagebox.showerror("Erro", "Nenhum DOCX selecionado.")
        return

    numero_whatsapp = entry_numero_whatsapp.get().strip()
    if not numero_whatsapp:
        messagebox.showerror("Erro", "Por favor, insira o número do WhatsApp.")
        return

    whatsapp_url = f"https://api.whatsapp.com/send?phone=55{numero_whatsapp}&text=Olá! Em anexo está o orçamento"

    # Encode para URL
    whatsapp_url_encoded = webbrowser.quote(whatsapp_url)

    # Abre a URL no navegador padrão
    webbrowser.open(whatsapp_url_encoded)

root = tk.Tk()
root.title("Sistema de Ordem de Serviço")

abas = ttk.Notebook(root)
aba_dados_e_servicos = tk.Frame(abas)
aba_enviar = tk.Frame(abas)  # Aba "Enviar"
abas.add(aba_dados_e_servicos, text='Dados e Serviços')
abas.add(aba_enviar, text='Enviar')
abas.pack(expand=1, fill='both')

# Diretório dos ícones
icone_pasta = Path('icons')

# Carregar e redimensionar os ícones
icone_salvar = carregar_icone('salvar.png', tamanho=(24, 24))
icone_abrir = carregar_icone('abrir.png', tamanho=(24, 24))
icone_abrir_numero = carregar_icone('abrir_numero.png', tamanho=(24, 24))
icone_whatsapp = carregar_icone('whatsapp.png', tamanho=(24, 24))

# Labels e Entradas na aba "Dados e Serviços"
tk.Label(aba_dados_e_servicos, text="Nome do Cliente:").grid(row=0, column=0, padx=2, pady=2, sticky='w')
tk.Label(aba_dados_e_servicos, text="Telefone:").grid(row=1, column=0, padx=2, pady=2, sticky='w')
tk.Label(aba_dados_e_servicos, text="Veículo:").grid(row=2, column=0, padx=2, pady=2, sticky='w')

entry_nome_cliente = tk.Entry(aba_dados_e_servicos, width=50)
entry_telefone = tk.Entry(aba_dados_e_servicos, width=50)
entry_veiculo = tk.Entry(aba_dados_e_servicos, width=50)

entry_nome_cliente.grid(row=0, column=1, padx=2, pady=2, sticky='w')
entry_telefone.grid(row=1, column=1, padx=2, pady=2, sticky='w')
entry_veiculo.grid(row=2, column=1, padx=2, pady=2, sticky='w')

tk.Label(aba_dados_e_servicos, text="").grid(row=4, column=0, padx=2, pady=5, columnspan=2)

# Frame para peças e valores
frame_pecas = tk.Frame(aba_dados_e_servicos, borderwidth=2, relief='groove')
frame_pecas.grid(row=5, column=0, columnspan=2, padx=5, pady=5, sticky='nsew')

# Labels para peças, quantidade, valor e total
tk.Label(frame_pecas, text="Peças").grid(row=0, column=0, padx=5, pady=5, sticky='w')
tk.Label(frame_pecas, text="QTD").grid(row=0, column=1, padx=5, pady=5, sticky='w')
tk.Label(frame_pecas, text="R$").grid(row=0, column=2, padx=5, pady=5, sticky='w')
tk.Label(frame_pecas, text="Total").grid(row=0, column=3, padx=5, pady=5, sticky='w')

num_linhas = 10
peca_entries = []
valor_entries = []
quantidade_entries = []
total_entries = []

# Loop para criar entradas para peças, valores, quantidades e totais
for i in range(num_linhas):
    peca_entry = tk.Entry(frame_pecas, width=30, justify='left')
    valor_entry = tk.Entry(frame_pecas, width=6, justify='left')
    quantidade_entry = tk.Entry(frame_pecas, width=6, justify='left')
    total_entry = tk.Entry(frame_pecas, width=15, state='readonly', justify='left')

    peca_entry.grid(row=i + 1, column=0, padx=5, pady=2, sticky='w')
    valor_entry.grid(row=i + 1, column=1, padx=5, pady=2, sticky='w')
    quantidade_entry.grid(row=i + 1, column=2, padx=5, pady=2, sticky='w')
    total_entry.grid(row=i + 1, column=3, padx=5, pady=2, sticky='w')

    valor_entry.bind('<KeyRelease>', atualizar_total)
    quantidade_entry.bind('<KeyRelease>', atualizar_total)

    peca_entries.append(peca_entry)
    valor_entries.append(valor_entry)
    quantidade_entries.append(quantidade_entry)
    total_entries.append(total_entry)

# Botões e Labels com ícones na aba "Dados e Serviços"
tk.Button(aba_dados_e_servicos, text="Salvar Ordem de Serviço", image=icone_salvar, command=salvar_ordem).grid(row=6, column=0, sticky='w', pady=0)
tk.Label(aba_dados_e_servicos, text="Salvar Ordem de Serviço").grid(row=6, column=1, sticky='w', pady=5)

# Botão para abrir o arquivo salvo
botao_abrir = tk.Button(aba_dados_e_servicos, text="Abrir Arquivo Salvo", image=icone_abrir, command=abrir_arquivo, state=tk.DISABLED)
botao_abrir.grid(row=6, column=2, sticky='w', padx=5, pady=5)
tk.Label(aba_dados_e_servicos, text="Abrir Arquivo Salvo").grid(row=6, column=3, sticky='w', padx=5, pady=5)

# Botão para abrir o arquivo pelo número
tk.Button(aba_dados_e_servicos, text="Abrir Arquivo pelo Número", image=icone_abrir_numero, command=abrir_arquivo_pelo_numero).grid(row=7, column=2, sticky='w', pady=5)
tk.Label(aba_dados_e_servicos, text="Abrir Arquivo pelo Número").grid(row=7, column=3, sticky='w', pady=5)

# Entrada para número da OS e botão para abrir arquivo pelo número
tk.Label(aba_dados_e_servicos, text="Abrir OS:").grid(row=7, column=0, padx=2, pady=5, sticky='w')
entry_numero_os = tk.Entry(aba_dados_e_servicos, width=10)
entry_numero_os.grid(row=7, column=1, padx=2, pady=2, sticky='w')

arquivo_salvo = [None]

# Label e entrada na aba "Enviar"
tk.Label(aba_enviar, text="Número do WhatsApp com DDD:").pack(pady=5)
entry_numero_whatsapp = tk.Entry(aba_enviar, width=50)
entry_numero_whatsapp.pack(pady=5)

# Botão para enviar mensagem pelo WhatsApp com o documento anexado
tk.Button(aba_enviar, text="Enviar Mensagem WhatsApp", image=icone_whatsapp, command=enviar_whatsapp).pack(pady=5)
tk.Label(aba_enviar, text="Enviar Mensagem WhatsApp").pack(pady=5)

# Label e lista de documentos DOCX na aba "Enviar"
tk.Label(aba_enviar, text="Selecionar DOCX:").pack(pady=5)
listbox_pdfs = tk.Listbox(aba_enviar, width=70, height=15)
listbox_pdfs.pack(pady=5)

# Atualiza a lista de documentos DOCX disponíveis
listbox_pdfs.bind('<Double-1>', abrir_arquivo_selecionado)
carregar_docx()

# Executa a interface gráfica
root.mainloop()
