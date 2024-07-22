import tkinter as tk
from tkinter import messagebox
from docx import Document
from pathlib import Path

def obter_numero_sequencial():
    """Obtem o próximo número sequencial para a OS."""
    arquivo_numero = Path('numero_sequencial.txt')
    
    if arquivo_numero.exists():
        with open(arquivo_numero, 'r') as file:
            numero = int(file.read().strip())
    else:
        numero = 0

    proximo_numero = numero + 1
    
    with open(arquivo_numero, 'w') as file:
        file.write(str(proximo_numero))
    
    return proximo_numero

def salvar_ordem():
    nome_cliente = entry_nome_cliente.get()
    endereco = entry_endereco.get()
    veiculo = entry_veiculo.get()
    servicos = text_servicos.get("1.0", tk.END).strip()

    if not nome_cliente or not endereco or not veiculo or not servicos:
        messagebox.showerror("Erro", "Todos os campos devem ser preenchidos!")
        return

    numero_sequencial = obter_numero_sequencial()
    nome_arquivo = f'Ordem_de_Servico_{numero_sequencial:05d}.docx'

    # Cria um novo documento do Word
    doc = Document()
    doc.add_heading('Ordem de Serviço', level=1)

    # Adiciona as informações ao documento
    doc.add_paragraph(f'Número da OS: {numero_sequencial:05d}')
    doc.add_paragraph(f'Nome do Cliente: {nome_cliente}')
    doc.add_paragraph(f'Endereço: {endereco}')
    doc.add_paragraph(f'Veículo: {veiculo}')
    doc.add_paragraph('Serviços:')
    doc.add_paragraph(servicos)

    # Salva o documento
    doc.save(nome_arquivo)
    messagebox.showinfo("Sucesso", f"Ordem de serviço salva com sucesso! ({nome_arquivo})")

# Configura a interface gráfica
root = tk.Tk()
root.title("Sistema de Ordem de Serviço")

tk.Label(root, text="Nome do Cliente:").grid(row=0, column=0, padx=10, pady=10)
tk.Label(root, text="Endereço:").grid(row=1, column=0, padx=10, pady=10)
tk.Label(root, text="Veículo:").grid(row=2, column=0, padx=10, pady=10)
tk.Label(root, text="Serviços:").grid(row=3, column=0, padx=10, pady=10)

entry_nome_cliente = tk.Entry(root, width=50)
entry_endereco = tk.Entry(root, width=50)
entry_veiculo = tk.Entry(root, width=50)

entry_nome_cliente.grid(row=0, column=1, padx=10, pady=10)
entry_endereco.grid(row=1, column=1, padx=10, pady=10)
entry_veiculo.grid(row=2, column=1, padx=10, pady=10)

# Adiciona um campo de texto para serviços
text_servicos = tk.Text(root, height=10, width=50)
text_servicos.grid(row=3, column=1, padx=10, pady=10)

tk.Button(root, text="Salvar Ordem de Serviço", command=salvar_ordem).grid(row=4, column=0, columnspan=2, pady=20)

root.mainloop()
