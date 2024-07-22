import tkinter as tk
from tkinter import messagebox
from tkinter import filedialog
from tkinter import ttk
from docx import Document
from pathlib import Path
import os

def obter_numero_sequencial():
    """Obtém o próximo número sequencial para a OS."""
    arquivo_numero = Path('numero_sequencial.txt')
    
    if arquivo_numero.exists():
        with open(arquivo_numero, 'r') as file:
            numero = int(file.read().strip())
    else:
        numero = 0

    proximo_numero = numero + 1
    
    with open(arquivo_numero, 'w') as file:
        file.write(str(proximo_numero))
    
    return proximo_numero

def substituir_marcadores(doc, substituicoes):
    """Substitui marcadores de posição no documento."""
    for par, valor in substituicoes.items():
        for paragrafo in doc.paragraphs:
            if par in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(par, valor)
                
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for par, valor in substituicoes.items():
                    if par in celula.text:
                        celula.text = celula.text.replace(par, valor)

def salvar_ordem():
    nome_cliente = entry_nome_cliente.get()
    telefone = entry_telefone.get()
    veiculo = entry_veiculo.get()
    
    pecas = [e.get() for e in peca_entries]
    valores = [v.get() for v in valor_entries]
    quantidades = [q.get() for q in quantidade_entries]

    numero_sequencial = obter_numero_sequencial()
    nome_arquivo = f'{numero_sequencial:05d}.docx'  # Nome do arquivo apenas com o número

    # Carrega o modelo de documento
    doc = Document('modelo.docx')

    # Substitui os marcadores de posição no modelo
    substituicoes = {
        '{{NUMERO_OS}}': f'{numero_sequencial:05d}',
        '{{NOME_CLIENTE}}': nome_cliente if nome_cliente else '',
        '{{TELEFONE}}': telefone if telefone else '',
        '{{VEICULO}}': veiculo if veiculo else '',
        '{{PECAS_VALORES}}': ''  # Inicialmente vazio, será preenchido com a tabela
    }

    substituir_marcadores(doc, substituicoes)

    # Adiciona uma tabela para peças, quantidades, valores e totais
    tabela = doc.add_table(rows=1, cols=4)
    tabela.style = 'Table Grid'
    
    # Define o cabeçalho da tabela
    hdr_cells = tabela.rows[0].cells
    hdr_cells[0].text = 'Peças'
    hdr_cells[1].text = 'Quantidade'
    hdr_cells[2].text = 'Valor Unitário'
    hdr_cells[3].text = 'Total'

    total_geral = 0.0
    
    for peca, valor, quantidade in zip(pecas, valores, quantidades):
        try:
            valor_float = float(valor.replace(',', '.')) if valor else 0.0
            quantidade_int = int(quantidade) if quantidade else 0
            total_item = valor_float * quantidade_int
            total_geral += total_item
        except ValueError:
            continue  # Ignora valores inválidos

        row_cells = tabela.add_row().cells
        row_cells[0].text = peca if peca else ''
        row_cells[1].text = f' {valor}' if valor else ''
        row_cells[2].text = quantidade if quantidade else ''
        row_cells[3].text = f' {total_item:.2f}'

    # Adiciona uma linha de resumo para quantidade total e soma dos valores
    try:
        quantidade_total = sum(int(q) for q in quantidades if q.isdigit())
    except ValueError:
        quantidade_total = 0

    row_cells = tabela.add_row().cells
    row_cells[0].text = 'Resumo'
    row_cells[1].text = ''
    row_cells[2].text = ''
    row_cells[3].text = f'R$ {total_geral:.2f}'

    # Substitui o marcador de posição da tabela
    substituicoes['{{PECAS_VALORES}}'] = tabela

    # Salva o documento
    doc.save(nome_arquivo)
    messagebox.showinfo("Sucesso", f"Ordem de serviço salva com sucesso! ({nome_arquivo})")
    
    # Habilita o botão de abrir o documento
    botao_abrir.config(state=tk.NORMAL)
    arquivo_salvo[0] = nome_arquivo  # Atualiza o nome do arquivo salvo

def abrir_arquivo():
    if arquivo_salvo[0]:
        os.startfile(arquivo_salvo[0])
    else:
        messagebox.showerror("Erro", "Nenhum arquivo salvo para abrir.")

def abrir_arquivo_pelo_numero():
    numero_os_str = entry_numero_os.get()  # Obtém o número da OS como string
    try:
        numero_os = int(numero_os_str)  # Converte a string para inteiro
        nome_arquivo = f'{numero_os:05d}.docx'  # Formata o nome do arquivo com zeros à esquerda
        if Path(nome_arquivo).exists():
            os.startfile(nome_arquivo)  # Abre o arquivo
        else:
            messagebox.showerror("Erro", "Arquivo não encontrado.")
    except ValueError:
        messagebox.showerror("Erro", "Número da OS inválido. Por favor, insira um número válido.")

def atualizar_total(event=None):
    total_geral = 0.0
    for i in range(num_linhas):
        try:
            valor = float(valor_entries[i].get().replace(',', '.')) if valor_entries[i].get() else 0.0
            quantidade = int(quantidade_entries[i].get()) if quantidade_entries[i].get() else 0
            total = valor * quantidade
            total_entries[i].config(state='normal')
            total_entries[i].delete(0, tk.END)
            total_entries[i].insert(0, f'R$ {total:.2f}')
            total_entries[i].config(state='readonly')
            total_geral += total
        except ValueError:
            continue  # Ignora valores inválidos
    return total_geral

# Configura a interface gráfica
root = tk.Tk()
root.title("Sistema de Ordem de Serviço")

# Abas
abas = ttk.Notebook(root)
aba_dados = tk.Frame(abas)
aba_servicos = tk.Frame(abas)
abas.add(aba_dados, text='Dados da Ordem de Serviço')
abas.add(aba_servicos, text='Peças e Valores')
abas.pack(expand=1, fill='both')

# Aba Dados da Ordem de Serviço
tk.Label(aba_dados, text="Nome do Cliente:").grid(row=0, column=0, padx=10, pady=10)
tk.Label(aba_dados, text="Telefone:").grid(row=1, column=0, padx=10, pady=10)
tk.Label(aba_dados, text="Veículo:").grid(row=2, column=0, padx=10, pady=10)

entry_nome_cliente = tk.Entry(aba_dados, width=50)
entry_telefone = tk.Entry(aba_dados, width=50)
entry_veiculo = tk.Entry(aba_dados, width=50)

entry_nome_cliente.grid(row=0, column=1, padx=10, pady=10)
entry_telefone.grid(row=1, column=1, padx=10, pady=10)
entry_veiculo.grid(row=2, column=1, padx=10, pady=10)

# Aba Peças e Valores
num_linhas = 10  # Número inicial de linhas preenchíveis
peca_entries = []
valor_entries = []
quantidade_entries = []
total_entries = []

# Cabeçalhos
tk.Label(aba_servicos, text="Peças").grid(row=0, column=0, padx=10, pady=10)
tk.Label(aba_servicos, text="QTD").grid(row=0, column=2, padx=10, pady=10)
tk.Label(aba_servicos, text="R$").grid(row=0, column=4, padx=10, pady=10)
tk.Label(aba_servicos, text="Total").grid(row=0, column=6, padx=10, pady=10)

# Linhas preenchíveis
for i in range(num_linhas):
    peca_entry = tk.Entry(aba_servicos, width=30)
    valor_entry = tk.Entry(aba_servicos, width=6)
    quantidade_entry = tk.Entry(aba_servicos, width=6)
    total_entry = tk.Entry(aba_servicos, width=15, state='readonly')

    peca_entry.grid(row=i + 1, column=0, padx=10, pady=5)
    valor_entry.grid(row=i + 1, column=2, padx=10, pady=5)
    quantidade_entry.grid(row=i + 1, column=4, padx=10, pady=5)
    total_entry.grid(row=i + 1, column=6, padx=10, pady=5)

    valor_entry.bind('<KeyRelease>', atualizar_total)
    quantidade_entry.bind('<KeyRelease>', atualizar_total)

    peca_entries.append(peca_entry)
    valor_entries.append(valor_entry)
    quantidade_entries.append(quantidade_entry)
    total_entries.append(total_entry)

# Botão para salvar a ordem de serviço
tk.Button(root, text="Salvar Ordem de Serviço", command=salvar_ordem).pack(pady=10)

# Campo para número da OS para abrir
tk.Label(root, text="Número da OS para abrir:").pack(pady=10)
entry_numero_os = tk.Entry(root, width=10)
entry_numero_os.pack(pady=5)

# Botão para abrir arquivo pelo número
tk.Button(root, text="Abrir Arquivo pelo Número", command=abrir_arquivo_pelo_numero).pack(pady=10)

# Botão para abrir o arquivo salvo
arquivo_salvo = [None]  # Lista para manter referência ao arquivo salvo
botao_abrir = tk.Button(root, text="Abrir Arquivo Salvo", command=abrir_arquivo, state=tk.DISABLED)
botao_abrir.pack(pady=10)

root.mainloop()
