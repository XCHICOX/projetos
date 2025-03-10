import tkinter as tk
from tkinter import messagebox
from tkinter import ttk
from docx import Document
from pathlib import Path

def obter_numero_sequencial():
    """Obtém o próximo número sequencial para a OS."""
    arquivo_numero = Path('numero_sequencial.txt')
    
    if arquivo_numero.exists():
        with open(arquivo_numero, 'r') as file:
            numero = int(file.read().strip())
    else:
        numero = 0

    proximo_numero = numero + 1
    
    with open(arquivo_numero, 'w') as file:
        file.write(str(proximo_numero))
    
    return proximo_numero

def substituir_marcadores(doc, substituicoes):
    """Substitui marcadores de posição no documento."""
    for par, valor in substituicoes.items():
        for paragrafo in doc.paragraphs:
            if par in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(par, valor)
                
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for par, valor in substituicoes.items():
                    if par in celula.text:
                        celula.text = celula.text.replace(par, valor)

def calcular_totais():
    """Calcula o total baseado nos valores e quantidades e atualiza a soma geral."""
    try:
        valores = [float(v.get().replace(',', '.')) for v in valor_entries if v.get().replace(',', '.').replace('.', '').isdigit()]
        quantidades = [int(q.get()) for q in quantidade_entries if q.get().isdigit()]
        totais = [v * q for v, q in zip(valores, quantidades)]
        total_geral = sum(totais)
        
        for total, total_entry in zip(totais, total_entries):
            total_entry.config(state=tk.NORMAL)
            total_entry.delete(0, tk.END)
            total_entry.insert(0, f'R$ {total:.2f}')
            total_entry.config(state=tk.DISABLED)
        
        total_geral_entry.config(state=tk.NORMAL)
        total_geral_entry.delete(0, tk.END)
        total_geral_entry.insert(0, f'R$ {total_geral:.2f}')
        total_geral_entry.config(state=tk.DISABLED)
        
    except ValueError:
        messagebox.showerror("Erro", "Valor ou quantidade inválido!")

def salvar_ordem():
    nome_cliente = entry_nome_cliente.get()
    telefone = entry_telefone.get()
    veiculo = entry_veiculo.get()
    
    pecas = [e.get() for e in peca_entries]
    valores = [v.get() for v in valor_entries]
    quantidades = [q.get() for q in quantidade_entries]

    numero_sequencial = obter_numero_sequencial()
    nome_arquivo = f'Ordem_de_Servico_{numero_sequencial:05d}.docx'

    # Carrega o modelo de documento
    doc = Document('modelo.docx')

    # Substitui os marcadores de posição no modelo
    substituicoes = {
        '{{NUMERO_OS}}': f'{numero_sequencial:05d}',
        '{{NOME_CLIENTE}}': nome_cliente if nome_cliente else '',
        '{{TELEFONE}}': telefone if telefone else '',
        '{{VEICULO}}': veiculo if veiculo else '',
        '{{PECAS_VALORES}}': ''  # Inicialmente vazio, será preenchido com a tabela
    }

    substituir_marcadores(doc, substituicoes)

    # Adiciona uma tabela para peças, valores, quantidades e totais
    tabela = doc.add_table(rows=1, cols=4)
    tabela.style = 'Table Grid'
    
    # Define o cabeçalho da tabela
    hdr_cells = tabela.rows[0].cells
    hdr_cells[0].text = 'Peças'
    hdr_cells[1].text = 'Valor Unitário'
    hdr_cells[2].text = 'Quantidade'
    hdr_cells[3].text = 'Total'

    total_geral = 0.0
    
    for peca, valor, quantidade in zip(pecas, valores, quantidades):
        try:
            valor_float = float(valor.replace(',', '.')) if valor else 0.0
            quantidade_int = int(quantidade) if quantidade else 0
            total_item = valor_float * quantidade_int
            total_geral += total_item
        except ValueError:
            continue  # Ignora valores inválidos

        row_cells = tabela.add_row().cells
        row_cells[0].text = peca if peca else ''
        row_cells[1].text = f'R$ {valor}' if valor else ''
        row_cells[2].text = quantidade if quantidade else ''
        row_cells[3].text = f'R$ {total_item:.2f}'

    # Adiciona uma linha de resumo para quantidade total e soma dos valores
    try:
        quantidade_total = sum(int(q) for q in quantidades if q.isdigit())
    except ValueError:
        quantidade_total = 0

    row_cells = tabela.add_row().cells
    row_cells[0].text = 'Resumo'
    row_cells[1].text = ''
    row_cells[2].text = f'Quantidade Total: {quantidade_total}'
    row_cells[3].text = f'Soma dos Valores: R$ {total_geral:.2f}'

    # Substitui o marcador de posição da tabela
    substituicoes['{{PECAS_VALORES}}'] = tabela

    # Salva o documento
    doc.save(nome_arquivo)
    messagebox.showinfo("Sucesso", f"Ordem de serviço salva com sucesso! ({nome_arquivo})")

# Configura a interface gráfica
root = tk.Tk()
root.title("Sistema de Ordem de Serviço")

# Configura o notebook (abas)
notebook = ttk.Notebook(root)
notebook.grid(row=0, column=0, padx=10, pady=10, sticky='nsew')

# Aba 1 - Dados do Cliente e Veículo
aba_dados = ttk.Frame(notebook)
notebook.add(aba_dados, text='Dados do Cliente')

tk.Label(aba_dados, text="Nome do Cliente:").grid(row=0, column=0, padx=10, pady=10)
tk.Label(aba_dados, text="Telefone:").grid(row=1, column=0, padx=10, pady=10)
tk.Label(aba_dados, text="Veículo:").grid(row=2, column=0, padx=10, pady=10)

entry_nome_cliente = tk.Entry(aba_dados, width=50)
entry_telefone = tk.Entry(aba_dados, width=50)
entry_veiculo = tk.Entry(aba_dados, width=50)

entry_nome_cliente.grid(row=0, column=1, padx=10, pady=10)
entry_telefone.grid(row=1, column=1, padx=10, pady=10)
entry_veiculo.grid(row=2, column=1, padx=10, pady=10)

# Aba 2 - Peças, Valores e Quantidade
aba_servicos = ttk.Frame(notebook)
notebook.add(aba_servicos, text='Peças e Valores')

tk.Label(aba_servicos, text="Peças:").grid(row=0, column=0, padx=10, pady=10)
tk.Label(aba_servicos, text="Valores:").grid(row=0, column=2, padx=10, pady=10)
tk.Label(aba_servicos, text="Quantidade:").grid(row=0, column=4, padx=10, pady=10)
tk.Label(aba_servicos, text="Total:").grid(row=0, column=6, padx=10, pady=10)

peca_entries = []
valor_entries = []
quantidade_entries = []
total_entries = []

# Aumenta a quantidade de linhas preenchíveis
num_linhas = 10  # Ajuste o número de linhas preenchíveis conforme necessário

for i in range(num_linhas):  # Ajuste o número conforme necessário
    peca_entry = tk.Entry(aba_servicos, width=40)  # Aumente o width aqui
    peca_entry.grid(row=i+1, column=0, padx=10, pady=5)
    peca_entries.append(peca_entry)

    valor_entry = tk.Entry(aba_servicos, width=30)  # Aumente o width aqui
    valor_entry.grid(row=i+1, column=2, padx=10, pady=5)
    valor_entries.append(valor_entry)

    quantidade_entry = tk.Entry(aba_servicos, width=15)
    quantidade_entry.grid(row=i+1, column=4, padx=10, pady=5)
    quantidade_entries.append(quantidade_entry)

    total_entry = tk.Entry(aba_servicos, width=20, state=tk.DISABLED)
    total_entry.grid(row=i+1, column=6, padx=10, pady=5)
    total_entries.append(total_entry)

tk.Label(aba_servicos, text="Total Geral:").grid(row=num_linhas+1, column=5, padx=10, pady=10)
total_geral_entry = tk.Entry(aba_servicos, width=20, state=tk.DISABLED)
total_geral_entry.grid(row=num_linhas+1, column=6, padx=10, pady=10)

# Adiciona a funcionalidade de cálculo
for quantidade_entry in quantidade_entries:
    quantidade_entry.bind('<KeyRelease>', lambda event: calcular_totais())

for valor_entry in valor_entries:
    valor_entry.bind('<KeyRelease>', lambda event: calcular_totais())

tk.Button(root, text="Salvar Ordem de Serviço", command=salvar_ordem).grid(row=1, column=0, columnspan=7, pady=20)

root.mainloop()
