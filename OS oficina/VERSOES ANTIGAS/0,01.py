import tkinter as tk
from tkinter import messagebox
from docx import Document
from docx.shared import Pt

def salvar_ordem():
    nome_cliente = entry_nome_cliente.get()
    endereco = entry_endereco.get()
    veiculo = entry_veiculo.get()

    if not nome_cliente or not endereco or not veiculo:
        messagebox.showerror("Erro", "Todos os campos devem ser preenchidos!")
        return

    # Cria um novo documento do Word
    doc = Document()
    doc.add_heading('Ordem de Serviço', level=1)

    # Adiciona as informações ao documento
    doc.add_paragraph(f'Nome do Cliente: {nome_cliente}')
    doc.add_paragraph(f'Endereço: {endereco}')
    doc.add_paragraph(f'Veículo: {veiculo}')

    # Salva o documento
    doc.save('Ordem_de_Servico.docx')
    messagebox.showinfo("Sucesso", "Ordem de serviço salva com sucesso!")

# Configura a interface gráfica
root = tk.Tk()
root.title("Sistema de Ordem de Serviço")

tk.Label(root, text="Nome do Cliente:").grid(row=0, column=0, padx=10, pady=10)
tk.Label(root, text="Endereço:").grid(row=1, column=0, padx=10, pady=10)
tk.Label(root, text="Veículo:").grid(row=2, column=0, padx=10, pady=10)

entry_nome_cliente = tk.Entry(root, width=50)
entry_endereco = tk.Entry(root, width=50)
entry_veiculo = tk.Entry(root, width=50)

entry_nome_cliente.grid(row=0, column=1, padx=10, pady=10)
entry_endereco.grid(row=1, column=1, padx=10, pady=10)
entry_veiculo.grid(row=2, column=1, padx=10, pady=10)

tk.Button(root, text="Salvar Ordem de Serviço", command=salvar_ordem).grid(row=3, column=0, columnspan=2, pady=20)

root.mainloop()
