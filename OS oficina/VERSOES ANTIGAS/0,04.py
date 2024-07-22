import tkinter as tk
from tkinter import messagebox
from docx import Document
from pathlib import Path

def obter_numero_sequencial():
    """Obtém o próximo número sequencial para a OS."""
    arquivo_numero = Path('numero_sequencial.txt')
    
    if arquivo_numero.exists():
        with open(arquivo_numero, 'r') as file:
            numero = int(file.read().strip())
    else:
        numero = 0

    proximo_numero = numero + 1
    
    with open(arquivo_numero, 'w') as file:
        file.write(str(proximo_numero))
    
    return proximo_numero

def substituir_marcadores(doc, substituicoes):
    """Substitui marcadores de posição no documento."""
    for par, valor in substituicoes.items():
        for paragrafo in doc.paragraphs:
            if par in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(par, valor)
                
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for par, valor in substituicoes.items():
                    if par in celula.text:
                        celula.text = celula.text.replace(par, valor)

def salvar_ordem():
    nome_cliente = entry_nome_cliente.get()
    telefone = entry_telefone.get()
    veiculo = entry_veiculo.get()
    servicos = text_servicos.get("1.0", tk.END).strip()
    valores = text_valores.get("1.0", tk.END).strip()

    if not nome_cliente or not telefone or not veiculo or not servicos or not valores:
        messagebox.showerror("Erro", "Todos os campos devem ser preenchidos!")
        return

    if len(servicos.split('\n')) != len(valores.split('\n')):
        messagebox.showerror("Erro", "O número de serviços deve corresponder ao número de valores!")
        return

    numero_sequencial = obter_numero_sequencial()
    nome_arquivo = f'Ordem_de_Servico_{numero_sequencial:05d}.docx'

    # Carrega o modelo de documento
    doc = Document('modelo.docx')

    # Substitui os marcadores de posição no modelo
    substituicoes = {
        '{{NUMERO_OS}}': f'{numero_sequencial:05d}',
        '{{NOME_CLIENTE}}': nome_cliente,
        '{{TELEFONE}}': telefone,
        '{{VEICULO}}': veiculo,
        '{{SERVIÇOS_VALORES}}': ''  # Inicialmente vazio, será preenchido com a tabela
    }

    substituir_marcadores(doc, substituicoes)

    # Adiciona uma tabela para serviços e valores
    tabela = doc.add_table(rows=1, cols=2)
    tabela.style = 'Table Grid'
    
    # Define o cabeçalho da tabela
    hdr_cells = tabela.rows[0].cells
    hdr_cells[0].text = 'Serviço'
    hdr_cells[1].text = 'Valor'

    servicos_lista = servicos.split('\n')
    valores_lista = valores.split('\n')

    for servico, valor in zip(servicos_lista, valores_lista):
        row_cells = tabela.add_row().cells
        row_cells[0].text = servico
        row_cells[1].text = valor

    # Substitui o marcador de posição da tabela
    substituicoes['{{SERVIÇOS_VALORES}}'] = tabela

    # Salva o documento
    doc.save(nome_arquivo)
    messagebox.showinfo("Sucesso", f"Ordem de serviço salva com sucesso! ({nome_arquivo})")

# Configura a interface gráfica
root = tk.Tk()
root.title("Sistema de Ordem de Serviço")

tk.Label(root, text="Nome do Cliente:").grid(row=0, column=0, padx=10, pady=10)
tk.Label(root, text="Telefone:").grid(row=1, column=0, padx=10, pady=10)
tk.Label(root, text="Veículo:").grid(row=2, column=0, padx=10, pady=10)
tk.Label(root, text="Serviços:").grid(row=3, column=0, padx=10, pady=10)
tk.Label(root, text="Valores:").grid(row=3, column=2, padx=10, pady=10)

entry_nome_cliente = tk.Entry(root, width=50)
entry_telefone = tk.Entry(root, width=50)
entry_veiculo = tk.Entry(root, width=50)

entry_nome_cliente.grid(row=0, column=1, padx=10, pady=10)
entry_telefone.grid(row=1, column=1, padx=10, pady=10)
entry_veiculo.grid(row=2, column=1, padx=10, pady=10)

# Adiciona campos de texto para serviços e valores
text_servicos = tk.Text(root, height=10, width=30)
text_servicos.grid(row=3, column=1, padx=10, pady=10)

text_valores = tk.Text(root, height=10, width=20)
text_valores.grid(row=3, column=3, padx=10, pady=10)

tk.Button(root, text="Salvar Ordem de Serviço", command=salvar_ordem).grid(row=4, column=0, columnspan=4, pady=20)

root.mainloop()
